import tkinter as tk
from tkinter import filedialog
import tkinter.messagebox as messagebox
import openpyxl
import requests
import os
import docx
import PyPDF2
import re

def indirilecek_klasor_yolu_sec():
    dosya_yolu = filedialog.askdirectory()
    indirilecek_klasor_yolu_entry.delete(0, tk.END)
    indirilecek_klasor_yolu_entry.insert(0, dosya_yolu) 

def select_excel_file():
    # Excel dosyasını seçmek için dosya iletişim kutusunu açar
    file_path = filedialog.askopenfilename(filetypes=[("Excel Files", "*.xlsx")])
    # Metin kutusunun içeriğini temizler
    excel_file_entry.delete(0, tk.END)
    # Seçilen dosya yolunu metin kutusuna ekler
    excel_file_entry.insert(0, file_path)

def count_words_in_docx(file_path):
    # .docx dosyasındaki kelime ve paragraf sayısını hesaplar
    doc = docx.Document(file_path)
    word_count = 0
    paragraph_count = 0

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            paragraph_count += 1
            words = re.findall(r'\b\w+\b', paragraph.text)
            word_count += len(words)

    return word_count, paragraph_count

def count_words_in_pdf(file_path):
    # .pdf dosyasındaki kelime ve sayfa sayısını hesaplar
    with open(file_path, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        word_count = 0
        page_count = len(reader.pages)

        for page in reader.pages:
            text = page.extract_text()
            if text:
                words = re.findall(r'\b\w+\b', text)
                word_count += len(words)

    return word_count, page_count

def download_file(url, download_directory):
    # Verilen URL'den dosyayı indirir ve indirilen dosyanın yolunu döndürür
    response = requests.get(url)
    file_name = url.split('/')[-1]
    file_path = os.path.join(download_directory, file_name)

    with open(file_path, 'wb') as file:
        file.write(response.content)

    return file_path

def process_excel_file():
    # Excel dosyasını işler
    file_path = excel_file_entry.get()
    file_directory = os.path.dirname(file_path)
    download_directory = file_directory
    file_name = os.path.splitext(os.path.basename(file_path))[0]
    result_file_directory = os.path.dirname(file_path)  # Sonuç dosyasının kaydedileceği klasör
    if not os.path.exists(download_directory):
        os.makedirs(download_directory)

    try:
        # Excel dosyasını yükler
        workbook = openpyxl.load_workbook(file_path)
        worksheet = workbook.active
        total_files = sum(1 for row in worksheet.iter_rows(min_row=2) if row[0].value)  # satır sayar

        processed_files = 0
        for row in worksheet.iter_rows(min_row=2):  # Başlık satırını atlamak için min_row=2
            url = row[0].value
            if url:
                try:
                    # URL'den dosyayı indirir
                    file_path = download_file(url, download_directory)

                    if file_path.lower().endswith('.docx'):
                        # .docx dosyasıysa kelime ve paragraf sayısını hesaplar
                        word_count, paragraph_count = count_words_in_docx(file_path)
                        row[1].value = word_count
                        row[2].value = paragraph_count
                    elif file_path.lower().endswith('.pdf'):
                        # .pdf dosyasıysa kelime ve sayfa sayısını hesaplar
                        word_count, page_count = count_words_in_pdf(file_path)
                        row[1].value = word_count
                        row[2].value = page_count
                    else:
                        # Desteklenmeyen dosya formatı uyarısı
                        messagebox.showwarning("Uyarı", f"Desteklenmeyen dosya formatı: {file_path}")

                    # İndirilen dosyayı siler
                    os.remove(file_path)
                except Exception as e:
                    # İndirme hatası uyarısı
                    messagebox.showerror("Hata", f"{url} indirilirken bir hata oluştu:\n{str(e)}")
                
                processed_files += 1
                # İşlenen dosya sayısını günceller ve uygulama ekranını günceller
                process_label.config(text=f"Toplam dosya: {total_files} | İşlenen dosya: {processed_files}")
                window.update()
                
        result_file_path = os.path.join(result_file_directory, os.path.splitext(os.path.basename(file_name))[0] + " (işlenmiş dosya).xlsx")
        workbook.save(result_file_path)
        # İşlem tamamlandı bilgisi
        messagebox.showinfo("Bilgi", f"İşlem tamamlandı. Sonuçlar {result_file_path} dosyasına kaydedildi.")
    except Exception as e:
        # Excel dosyası açma hatası
        messagebox.showerror("Hata", f"Excel dosyası açılırken bir hata oluştu:\n{str(e)}")


def count_words_in_docx(file_path):
    doc = docx.Document(file_path)
    word_count = 0
    paragraph_count = 0

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            paragraph_count += 1
            words = re.findall(r'\b\w+\b', paragraph.text)
            word_count += len(words)

    return word_count, paragraph_count

def count_words_in_pdf(file_path):
    with open(file_path, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        word_count = 0
        page_count = len(reader.pages)

        for page in reader.pages:
            text = page.extract_text()
            if text:
                words = re.findall(r'\b\w+\b', text)
                word_count += len(words)

    return word_count, page_count

def download_file(url, download_directory):
    response = requests.get(url)
    file_name = url.split('/')[-1]
    file_path = os.path.join(download_directory, file_name)

    with open(file_path, 'wb') as file:
        file.write(response.content)

    return file_path

def pdf_indir():
    """
    Belirtilen URL'den PDF dosyasını indiren ve indirilen dosyanın bilgilerini ekrana yazdıran fonksiyon.
    
    Args:
        url (str): İndirilecek PDF dosyasının URL'si.
        pdf_kayit_klasoru (str): İndirilen PDF dosyalarının kaydedileceği klasörün yolunu içeren metin.
    """
    # URL'leri giriş kutusundan alın ve bir listeye ayır
    urls = url_entry.get("1.0", tk.END).strip().split('\n')
    total_url_count = sum(1 for url in urls if url.strip())
    # İndirilecek PDF dosyalarının kaydedileceği klasörü alın
    pdf_kayit_klasoru = indirilecek_klasor_yolu_entry.get()
    
    # Klasörü oluştur (varsa zaten var olduğundan hata vermez)
    os.makedirs(pdf_kayit_klasoru, exist_ok=True)
    
    # Klasör kontrolü ve oluşturma işlemi tekrar yapılır
    if not os.path.exists(pdf_kayit_klasoru):
        os.makedirs(pdf_kayit_klasoru)
    
    # İndirilen dosyaların sayısını tutacak değişken
    indirilen_dosya_sayisi = 0
    anlik_sayma = 0
    
    # Her bir URL için işlem yap
    for url in urls:
        if url:
             
           #anlik_sayma_label.config(text=f"Toplam dosya: {total_url_count}")
            try:
                # Dosyayı indir ve indirilen dosyanın yolunu al
                download_file(url, pdf_kayit_klasoru)
                indirilen_dosya_sayisi += 1

            except Exception as e:
                messagebox.showerror("Hata", f"{url} indirilirken bir hata oluştu:\n{str(e)}")
            anlik_sayma += 1
                # İşlenen dosya sayısını günceller ve uygulama ekranını günceller
            anlik_sayma_label.config(text=f"Toplam dosya: {total_url_count} | İşlenen dosya: {anlik_sayma}")
            window.update()

                                    
    # Uyarı mesajı ve indirilen dosya sayısını ekrana yazdır
    messagebox.showinfo("Bilgi", f"{indirilen_dosya_sayisi} dosya indirildi.")


# Tkinter uygulama penceresi oluşturulur
window = tk.Tk()
window.title("Toplu indirme ve sayma")
window.geometry("500x400")

# Excel dosyası seçme bileşenleri oluşturulur
excel_file_frame = tk.Frame(window)
excel_file_frame.pack(pady=20)

excel_file_label = tk.Label(excel_file_frame, text="Excel Dosyası:")
excel_file_label.pack(side=tk.LEFT)

excel_file_entry = tk.Entry(excel_file_frame, width=50)
excel_file_entry.pack(side=tk.LEFT)

select_excel_file_button = tk.Button(excel_file_frame, text="Excel Dosyası Seç", command=select_excel_file)
select_excel_file_button.pack(side=tk.LEFT)

# Excel dosyasını işleme butonu oluşturulur
process_excel_file_button = tk.Button(window, text="Excel Dosyasını İşle", command=process_excel_file)
process_excel_file_button.pack()

# İşlenen dosya sayacı oluşturulur
process_label = tk.Label(window, text="Toplam dosya: 0 | İşlenen dosya: 0")
process_label.pack(pady=10)

nota_label = tk.Label(window, text="*********************************************************************************************")
nota_label.pack()
# Klasör seçme
indirilecek_klasor_yolu_label = tk.Label(window, text="İndirilecek dosyaların linklerini yazınız")
indirilecek_klasor_yolu_label.pack()

url_entry = tk.Text(window, height=5, width=50)
url_entry.pack()

indirilecek_klasor_yolu_label = tk.Label(window, text="\n Toplu Kayıt Klasörü:")
indirilecek_klasor_yolu_label.pack()

indirilecek_klasor_yolu_entry = tk.Entry(window, width=50)
indirilecek_klasor_yolu_entry.pack()

indirilecek_klasor_yolu_button = tk.Button(window, text="İndirilecek Klasör Yolu seç", command=indirilecek_klasor_yolu_sec)
indirilecek_klasor_yolu_button.pack()

toplu_dosya_indir_button = tk.Button(window, text="Linklerdeki dosyaları indir", command=pdf_indir)
toplu_dosya_indir_button.pack()

anlik_sayma_label = tk.Label(window, text="Toplam dosya: 0 | İşlenen dosya: 0")
anlik_sayma_label.pack(pady=10)

# Tkinter uygulaması çalıştırılır
window.mainloop()
